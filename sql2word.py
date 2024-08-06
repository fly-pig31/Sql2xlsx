import json

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from pgsql.mypgsql import schema_info_s, Conn

conn = Conn("data", "postgres", "pwd", "192.168.1.1", "5432", "xhz")

input_string = '''

'''

# 清除字符串中的空格和换行符，并按换行符分割成列表
tables = input_string.strip().split('\n')

# 使用集合去重，再转换为列表
tables = list(set(tables))


def write_to_word(result, filename):
    document = Document()

    # Define styles for table elements
    header_font_size = Pt(10)
    cell_font_size = Pt(10)

    # Loop through the result and create tables in the Word document
    for item in result:
        # print(f'item:{item}')
        table_name = item[0][0]
        table_comment = item[0][1]

        # Add an empty paragraph before each new table
        document.add_paragraph()

        # Create a new table for table name and comment
        table_header = document.add_table(rows=1, cols=1)
        table_header.width = Pt(500)
        table_header.style = 'Table Grid'

        # Add table name as a cell in the header table
        cell_name = table_header.cell(0, 0)
        cell_name.text = table_name if table_name else "N/A"  # Use "N/A" if table_name is None or empty
        cell_name.paragraphs[0].runs[0].font.size = header_font_size
        cell_name.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add table comment as a cell in the header table
        cell_comment = table_header.add_row().cells[0]
        cell_comment.text = table_comment.upper() if table_comment else "N/A"  # Use "N/A" if table_comment is None or empty
        cell_comment.paragraphs[0].runs[0].font.size = header_font_size
        cell_comment.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Create a new table for the actual data
        table_data = document.add_table(rows=1, cols=7)
        table_data.width = Pt(500)
        table_data.style = 'Table Grid'

        # Set header row style
        header_row = table_data.rows[0]
        header_row.cells[0].text = "字段名"
        header_row.cells[1].text = "注释"
        header_row.cells[2].text = "数据类型"
        header_row.cells[3].text = "长度"
        header_row.cells[4].text = "小数位数"
        header_row.cells[5].text = "是否允许为空"
        header_row.cells[6].text = "默认值"

        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].font.size = header_font_size
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Set cell style for the rest of the rows
        for row in item[1:]:
            new_row = table_data.add_row().cells
            for i, cell_data in enumerate(row[2:]):
                new_row[i].text = str(cell_data).upper()
                new_row[i].paragraphs[0].runs[0].font.size = cell_font_size
                new_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(filename)


line = '''
---------------------------------------------------------------------------
'''
result = schema_info_s(tables, conn)
# sorted_result = sorted(result, key=lambda x: x[0][1].split("_")[2])
sorted_result = sorted(result, key=lambda x: x[0][1].split("_")[2] if len(x[0][1].split("_")) >= 3 else '')
# for row in sorted_result:
#     print(row)
print(f'选中的表:{tables}')
# for item in result:
#     for row in item:
#         print(f"row:{row}")
#     print(line)
print("Executing write function")
filename="word_name"
docx_name = f"{filename}.docx"
xlsx_name = f"{filename}.docx"
write_to_word(sorted_result, docx_name)
