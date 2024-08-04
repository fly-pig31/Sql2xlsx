from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

"""
写入word文档
result: 表示查询的结果列表,通常是一个元组.
filename: 要写入的文件名
th_rows(table_header_rows): 要输出的结果的表头信息(字段名、数据类型、长度、小数位数、是否允许为空、默认值等等显示在表头的信息.同时列表的长度也表示从每条结果中读取的数据量)
"""


def write_to_word(result: tuple, filename: str, table_header_rows: list):
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # Define styles for table elements
    header_font_size = Pt(10)
    cell_font_size = Pt(10)

    # Loop through the result and create tables in the Word document
    for item in result:
        # print(f'item:{item}')
        table_name = item[0][0]
        table_comment = item[0][1]

        # Add an empty paragraph before each new table
        document.add_paragraph()

        # Create a new table for table name and comment
        table_header = document.add_table(rows=1, cols=1)
        table_header.width = Pt(500)
        table_header.style = 'Table Grid'

        # Add table name as a cell in the header table
        cell_name = table_header.cell(0, 0)
        cell_name.text = table_name if table_name else "表名"  # Use "表名" if table_name is None or empty
        cell_name.paragraphs[0].runs[0].font.size = header_font_size
        cell_name.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add table comment as a cell in the header table
        cell_comment = table_header.add_row().cells[0]
        cell_comment.text = table_comment.upper() if table_comment else "N/A"  # Use "N/A" if table_comment is None or empty
        cell_comment.paragraphs[0].runs[0].font.size = header_font_size
        cell_comment.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Create a new table for the actual data
        table_data = document.add_table(rows=1, cols=len(table_header_rows))
        table_data.width = Pt(500)
        table_data.style = 'Table Grid'

        # Set header row style
        header_row = table_data.rows[0]
        for i in table_header_rows:
            header_row.cells[table_header_rows.index(i)].text = i

        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].font.size = header_font_size
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Set cell style for the rest of the rows
        for row in item[1:]:
            new_row = table_data.add_row().cells
            for i, cell_data in enumerate(row[2:]):
                new_row[i].text = str(cell_data).upper()
                new_row[i].paragraphs[0].runs[0].font.size = cell_font_size
                new_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(filename)
    print("写入成功")

def write_to_excel(result, filename, table_header_rows: list):
    wb = Workbook()
    ws = wb.active
    # 设置表头样式
    header_font = Font(bold=True)
    header_fill = PatternFill("solid", fgColor="92d050")  # 绿色背景色
    data_alignment = Alignment(horizontal='center', vertical='center')
    # 设置边框样式
    border = Border(left=Side(border_style='thin', color='000000'),
                    right=Side(border_style='thin', color='000000'),
                    top=Side(border_style='thin', color='000000'),
                    bottom=Side(border_style='thin', color='000000'))
    row_num = 0
    row_num_list = []
    for index, item in enumerate(result):
        if index == 0:
            row_num = 1
        else:
            length = len(result[index - 1])
            row_num += length + 4
            # 写入表注释
        row_num_list.append(row_num)
        # 写入表注释和表名
        ws.merge_cells(f"A{row_num}:G{row_num}")
        ws.merge_cells(f"A{row_num + 1}:G{row_num + 1}")
        for i in range(2):
            table_name_cell = ws.cell(row=row_num + i, column=1, value=item[0][i])
            table_name_cell.font = header_font
            table_name_cell.alignment = data_alignment
            table_name_cell.fill = header_fill
        header = table_header_rows
        ws.append(header)
        for cell in ws[ws.max_row]:
            cell.font = header_font
            cell.fill = header_fill  # 设置填充颜色
            cell.alignment = data_alignment
        # 写入数据剩下的数据
        for row in item:
            ws.append(row[2:])
    for row in ws.iter_rows():
        for cell in row:
            if cell.row + 1 not in row_num_list:
                cell.border = border
    wb.save(filename)
    print("写入成功")
