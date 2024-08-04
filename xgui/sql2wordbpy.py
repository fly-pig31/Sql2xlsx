import queue
import threading
import time
import traceback
from tkinter import Tk, Button, messagebox, Label, Frame, Entry, IntVar, Radiobutton, StringVar, \
    constants, filedialog, Text, font, Scrollbar, BooleanVar, Checkbutton, ttk
from tkinter.font import Font
from tkinter.messagebox import showinfo, showerror
from tkinter.ttk import Progressbar
from enum import Enum
import logging
from pymysql import OperationalError
import psycopg2


class MsgEnum(Enum):
    """
        消息类型枚举
    """
    START = 0
    STOP = 1
    EXIT = 3


class InputValue:
    def __init__(self, host, port, user, password, database, sql, calc_mode, isAllTable):
        self.host = host
        self.port = port
        self.user = user
        self.password = password
        self.database = database
        self.sql = sql
        self.calc_mode = calc_mode
        self.isAllTable = isAllTable


class GuiTempldate:
    class Cache:
        RUNING = False

    def __init__(self) -> None:
        # 1、Gui消息队列
        self.DB = None
        self.database = None
        self.msg_center = MsgCenter(self)
        # 2、窗口设置
        self.root = Tk()
        self.root.title("数据库导出工具")
        # self.root.wm_resizable(False, False) # 设置窗口大小不可拉伸
        self.root.geometry('600x700+500+200')  # 设置窗口： 宽 x 高 + 窗口位置x坐标 + 窗口位置y坐标
        self.root.protocol("WM_DELETE_WINDOW", self.close_event)
        # 4、初始化各个组件和布局
        self.initGui()

    def initGui(self):
        # 1- 标签
        text_str = """版本: 1.0.0
        #author burukeyou
        #说明:
            1) 支持导出数据库所有表结构信息
            2）支持导出数据库表清单
            3）支持自定义sql导出
        """
        Label(self.root, text=text_str, justify='left', fg='red').pack(anchor=constants.W)
        fm02 = Frame(self.root)
        self.user_name = StringVar(value="root")
        fm02.pack(anchor=constants.W, fill=constants.X)
        Label(fm02, text='选择数据库 ').pack(side=constants.LEFT)
        # 下拉框
        self.database_options = ttk.Combobox(fm02, width=38, textvariable=self.DB)
        self.database_options.pack(side=constants.RIGHT)
        self.database_options['values'] = ['mysql', 'postgres', 'Oracle']

        # 4- 文本框
        fm02 = Frame(self.root)
        self.ip_address = StringVar(value="localhost:3306")
        fm02.pack(anchor=constants.W, fill=constants.X)
        Label(fm02, text='服务器地址 ').pack(side=constants.LEFT)
        self.name_btn = Entry(fm02, width=40, textvariable=self.ip_address)
        self.name_btn.pack(side=constants.RIGHT)

        # 4- 文本框
        fm02 = Frame(self.root)
        self.user_name = StringVar(value="root")
        fm02.pack(anchor=constants.W, fill=constants.X)
        Label(fm02, text='帐号').pack(side=constants.LEFT)
        self.name_btn = Entry(fm02, width=40, textvariable=self.user_name)
        self.name_btn.pack(side=constants.RIGHT)

        # 4- 文本框
        fm02 = Frame(self.root)
        self.password = StringVar(value="123456")
        fm02.pack(anchor=constants.W, fill=constants.X)
        Label(fm02, text='密码').pack(side=constants.LEFT)
        self.name_btn = Entry(fm02, width=40, textvariable=self.password)
        self.name_btn.pack(side=constants.RIGHT)

        # 4- 文本框
        fm02 = Frame(self.root)
        self.database = StringVar(value="")
        fm02.pack(anchor=constants.W, fill=constants.X)
        Label(fm02, text='数据库').pack(side=constants.LEFT)
        self.name_btn = Entry(fm02, width=40, textvariable=self.database)
        self.name_btn.pack(side=constants.RIGHT)

        # 3- 单选框
        self.mode_var = IntVar(value=1)
        fm01 = Frame(self.root)  # , bg='blue'
        fm01.pack(anchor=constants.W)
        Label(fm01, text='导出模式').grid(row=0, column=0, sticky='W')
        Radiobutton(fm01, text='表结构', variable=self.mode_var, value=1).grid(row=0, column=1, sticky='E', padx=40)
        Radiobutton(fm01, text='表清单', variable=self.mode_var, value=2).grid(row=0, column=2, sticky='E', padx=40)
        Radiobutton(fm01, text='自定义SQL', variable=self.mode_var, value=3).grid(row=0, column=3, sticky='E', padx=40)

        # 5- 勾选框
        self.is_all_flag = BooleanVar()
        self.is_all_table_btn = Checkbutton(self.root, variable=self.is_all_flag,
                                            text='所有表执行一遍(自定义SQL可使用模版变量#{tableName})', onvalue=True,
                                            offvalue=False)
        self.is_all_table_btn.pack(anchor=constants.W)

        # 9- 多行文本框
        Label(self.root, text='自定义SQL输入').pack(anchor=constants.W)
        fm22 = Frame(self.root)
        fm22.pack(anchor=constants.W, fill=constants.X)
        scroll = Scrollbar(fm22)
        scroll.pack(side=constants.RIGHT, fill=constants.Y)
        ft = Font(family='宋体', size=14, weight=font.BOLD)
        self.text_btn = Text(fm22, height=12, fg="black", font=ft, bg="white", insertbackground="black")
        self.text_btn.pack(side=constants.LEFT)

        # text 联动 scroll
        scroll.config(command=self.text_btn.yview)
        self.text_btn.config(yscrollcommand=scroll.set)

        # 进度条
        self.progressbar = Progressbar(self.root, length=600, mode="determinate", orient=constants.HORIZONTAL)
        self.progressbar.pack(anchor=constants.W)

        self.progressbar['maximum'] = 100
        # 10、开始结束按钮
        fm06 = Frame(self.root)
        fm06.pack(side=constants.BOTTOM)
        self.start_btn = Button(fm06, text="导出", width=6, height=1, command=self.start_event)
        self.start_btn.grid(row=0, column=0, sticky='W', padx=20, pady=20)
        self.stop_btn = Button(fm06, text="停止", width=6, height=1, command=self.stop_event)
        self.stop_btn.grid(row=0, column=1, sticky='E', padx=20, pady=20)
        self.test_btn = Button(fm06, text="测试连接", width=6, height=1, command=self.test_connection)
        self.test_btn.grid(row=0, column=2, sticky='E', padx=20, pady=20)

    def start_event(self):
        self.msg_center.put(MsgEnum.START)
        self.Cache.RUNING = True
        self.main_calc_thread: threading.Thread = threading.Thread(target=start_clac, args=(self,))
        self.main_calc_thread.start()

    def stop_event(self):
        self.msg_center.put(MsgEnum.STOP)
        self.Cache.RUNING = False

    def close_event(self):
        # 关闭窗口
        if self.Cache.RUNING and not messagebox.askokcancel("警告", "任务还在执行中，确定要关闭吗?"):
            return
        self.root.destroy()
        self.msg_center.put(MsgEnum.EXIT)

    def test_connection(self) -> bool:
        input_value = self.getInputValue()
        if not checkParam(input_value):
            return False
        conn = None
        try:
            conn = common_connect(self.DB, input_value)
            showinfo(message="测试连接成功")
            return True

        except OperationalError as e:
            showerror(message="测试连接数据库失败,请确认数据库信息是否正确。 \n" + str(e))
        finally:
            if conn:
                conn.close()

    def showUI(self):
        # 启动消息队列
        threading.Thread(target=self.msg_center.mainloop).start()
        # 这个方法会循环阻塞住，监听gui的事件,得放到主线程最后
        self.root.mainloop()

    def getInputValue(self) -> InputValue:
        ipaddress: str = self.ip_address.get().strip()
        if ipaddress == "" or ipaddress.find(":") < 0:
            showerror(message="请输入正确的ip地址")
            return
        ip = ipaddress.split(":")[0].strip()
        port = ipaddress.split(":")[1].strip()
        database = self.database.get().strip()
        user_name = self.user_name.get().strip()
        password = self.password.get().strip()
        calc_mode = self.mode_var.get()
        custom_sql = self.text_btn.get("1.0", "end-1c").strip()
        isAllTable: bool = self.is_all_flag.get()
        return InputValue(ip, port, user_name, password, database, custom_sql, calc_mode, isAllTable)


def start_clac(gui: GuiTempldate):
    try:
        # 实际计算任务
        start_demo(gui)
    finally:
        # 发布结束事件
        gui.stop_event()


class MsgCenter:
    """
        消息队列
            主要处理窗口控件消息
    """

    def __init__(self, obj: GuiTempldate) -> None:
        self.queue = queue.Queue()
        self.obj = obj

    def put(self, msg: Enum):
        self.queue.put(msg)

    def mainloop(self):
        while True:
            try:
                # 阻塞获取消息
                msg = self.queue.get()
                print("消费消息: {}".format(msg))
                if msg == MsgEnum.START:
                    MsgStrategy.start_strategy(self.obj)
                elif msg == MsgEnum.STOP:
                    MsgStrategy.stop_strategy(self.obj)
                elif msg == MsgEnum.EXIT:
                    break
                else:
                    pass
            except queue.Empty:
                traceback.print_exc()


class MsgStrategy:
    @classmethod
    def start_strategy(cls, gui: GuiTempldate):
        gui.start_btn.config(state=constants.DISABLED)
        gui.stop_btn.config(state=constants.NORMAL)
        #
        gui.progressbar['value'] = 0
        for i in range(100):
            time.sleep(0.1)
            gui.progressbar['value'] += 1
            if i >= 50:
                gui.progressbar['value'] = 100
                gui.root.update()
                break
            gui.root.update()

    @classmethod
    def stop_strategy(cls, gui: GuiTempldate):
        gui.start_btn.config(state=constants.NORMAL)
        gui.stop_btn.config(state=constants.DISABLED)


# ===============================业务逻辑 ======================================================
from typing import List
import pymysql
import pandas as pd
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
from docx.table import Table, _Cell


def checkParam(input_value: InputValue) -> bool:
    if input_value.database == "" or input_value.user == "" or input_value.password == "" or input_value.host == "" or input_value.port == "":
        showerror(message="请填写完整的数据库连接参数")
        return False
    return True


# def connectDB(input_value: InputValue):
#     return pymysql.connect(host=input_value.host,
#                            user=input_value.user,
#                            password=input_value.password,
#                            db=input_value.database,
#                            port=int(input_value.port), charset="utf8")
#
#
# def connect_postgres(input_value: InputValue):
#     return psycopg2.connect(host=input_value.host,
#                             user=input_value.user,
#                             password=input_value.password,
#                             database=input_value.database,
#                             port=int(input_value.port), charset="utf8")


def common_connect(db: str, input_value: InputValue):
    if db == 'mysql':
        return pymysql.connect(host=input_value.host,
                               user=input_value.user,
                               password=input_value.password,
                               db=input_value.database,
                               port=int(input_value.port), charset="utf8")
    if db == 'postgres':
        db = str(input_value.database).split('.')[0]
        return psycopg2.connect(host=input_value.host,
                                user=input_value.user,
                                password=input_value.password,
                                database=input_value.database,
                                port=int(input_value.port), charset="utf8")


def start_demo(gui: GuiTempldate):
    """ 自定义计算任务
    :param gui:             gui组件对象
    :return:
    """
    input_value = gui.getInputValue()
    # 校验参数
    if not checkParam:
        return
    conn = None
    try:
        # conn: pymysql.Connection = connectDB(input_value)
        conn = common_connect(gui.DB, input_value)
        print("建立连接成功")
        download(conn, input_value,db=gui.DB)
    except OperationalError as e:
        logging.error(e)
        showerror(message="连接数据库失败,请确认数据库信息是否正确")
    except Exception as e1:
        # 打印堆栈信息
        e1.with_traceback(None)
        showerror(message="导出失败")
    finally:
        if conn is not None:
            conn.close()
            print("关闭连接成功")


"""
    表结构SQL
"""


def tableSchemaSql(table_name, dataBase):
    return f"""
   select ORDINAL_POSITION 序号,
    COLUMN_NAME 字段名,
    COLUMN_TYPE 字段类型,
    -- DATA_TYPE 字段类型,
    -- CHARACTER_MAXIMUM_LENGTH 长度,
    -- IS_NULLABLE 是否为空,
    case when COLUMN_DEFAULT is null then '无' else COLUMN_DEFAULT end 默认值,
    COLUMN_COMMENT 字段说明
    FROM
    INFORMATION_SCHEMA.COLUMNS
    where
    table_schema ='{dataBase}'
    and table_name = '{table_name}';
   """


"""
    Postgres表结构SQL
"""


def postgres_table_schema_sql(table_name, dataBase):
    return f'''
                SELECT
                    tab.table_comment AS "表注释",
                    '{dataBase}.{table_name}' AS "表名",
                    cols.column_name AS "字段名",
                    pgd.description AS "注释",
                    cols.data_type AS "数据类型",
                    cols.character_maximum_length AS "长度",
                    cols.numeric_scale AS "小数位数",
                    cols.is_nullable AS "是否允许为空",
                    cols.column_default AS "默认值"
                FROM
                    information_schema.columns AS cols
                LEFT JOIN
                    pg_catalog.pg_description AS pgd ON cols.table_name::regclass = pgd.objoid
                        AND pgd.objsubid = cols.ordinal_position
                LEFT JOIN
                    (
                        SELECT
                            pgc.oid AS table_oid,
                            pd.description AS table_comment
                        FROM
                            pg_catalog.pg_class AS pgc
                        LEFT JOIN
                            pg_catalog.pg_description AS pd ON pgc.oid = pd.objoid
                            AND pd.objsubid = 0
                        WHERE
                            pgc.relname = '{dataBase}.{table_name}'
                    ) AS tab ON cols.table_name::regclass = tab.table_oid
                WHERE
                    cols.table_name = '{dataBase}.{table_name}'
                '''


"""
    数据库表清单SQL
"""


def databaseTableListSql(database):
    return f"""select table_name as '表名',TABLE_COMMENT '表中文名称','业务表',TABLE_COMMENT '说明'
               from information_schema.tables
               where table_schema='{database}' and  LOWER(table_type) ='base table'
    """


def readAllTableToDataFrame(conn, dataBase, db) -> dict:
    # 获取所有表的信息
    print("获取所有表的信息")
    tables = pd.read_sql_query("SHOW TABLES;", conn)
    df_map = {}
    for _, table in tables.iterrows():
        table_name = table[0]
        if table_name is None or table_name == "":
            continue
        # sql = tableSchemaSql(table_name, dataBase)
        sql = tableSchemaSql(table_name,
                             dataBase) if db == 'mysql' else None if db == 'postgres' else postgres_table_schema_sql(
            table_name,
            str(dataBase).split(
                '.')[1])
        df = pd.read_sql_query(sql, conn)
        df = df.reset_index(drop=True)
        df.index = df.index + 1
        df_map[table_name] = df
    return df_map


def readDataBaseTableListToDataFrame(conn, dataBase):
    df_map = {}
    sql = databaseTableListSql(dataBase)
    df = pd.read_sql_query(sql, conn)
    df = df.reset_index(drop=True)
    df.index = df.index + 1
    df_map["tmp"] = df
    return df_map


def readByCustomSqlToDataFrame(conn, input_value):
    # sql模版
    sql = input_value.sql
    # 获取所有表的信息
    tables = pd.read_sql_query("SHOW TABLES;", conn)
    df_map = {}
    for _, table in tables.iterrows():
        table_name = table[0]
        if table_name is None or table_name == "":
            continue
        # 替换变量名 #{taleName}
        tmp_sql = sql.replace("#{tableName}", table_name)
        df = pd.read_sql_query(tmp_sql, conn)
        df = df.reset_index(drop=True)
        df.index = df.index + 1
        df_map[table_name] = df
    return df_map


def downloadToWord(df_map, save_path, heading):
    # 创建Word文档
    doc = Document()
    for table_name in df_map:
        df = df_map[table_name]
        # 添加标题
        if heading:
            doc.add_heading(heading, level=3)
        else:
            doc.add_heading(f"【{table_name}表】", level=3)
        # 转成list表格数据
        table_data: List = [df.columns.tolist()] + df.values.tolist()
        # 创建 rows x col 表格
        table: Table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        # 设置表格样式
        table.style = "Table Grid"
        # 设置表格字体和大小
        # table.style.font.name = "宋体"
        # table.style.font.size = 5
        # 填充表格内容
        for i, row in enumerate(table_data):
            for j, value in enumerate(row):
                cell: _Cell = table.cell(i, j)
                cell.text = str(value)
                # 设置第一行
                if i == 0:
                    # 设置文本加粗
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        #  设置第一行背景色为浅灰色D9D9D9
        for i, cell in enumerate(table.rows[0].cells):
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm_1)
    # 保存Word文档
    doc.save(save_path)
    print("导出" + save_path + "成功")


def download(conn, input_value, db):
    dataBase = input_value.database
    calc_mode = input_value.calc_mode
    custom_sql = input_value.sql
    # 读取所有表结构到DataFrame
    heading = None
    df_map = {}
    if calc_mode == 1:
        # 所有表结构
        print("执行模式1")
        df_map = readAllTableToDataFrame(conn, dataBase, db)
    elif calc_mode == 2:
        # 表清单
        print("执行模式2")
        df_map = readDataBaseTableListToDataFrame(conn, dataBase)
        heading = f"【{dataBase}数据库表清单】"
    else:
        # 自定义sql
        if custom_sql is None or custom_sql == "":
            showinfo("", "请输入自定义SQL")
            return
        # 配置了都执行一遍并且sql模版真中包含#{tableName}变量
        if input_value.isAllTable == True and custom_sql.find("#{tableName}") != -1:
            print("自定义SQL对所有表执行一遍")
            df_map = readByCustomSqlToDataFrame(conn, input_value)
        else:
            df = pd.read_sql_query(custom_sql, conn)
            df_map["tmp"] = df
            heading = f"【{dataBase}自定义SQL数据】"
    # 选择导出位置
    now = datetime.now().strftime("%y%m%d%H%M%S")
    file_name = dataBase + f"数据库数据-{now}.docx"
    save_path = filedialog.asksaveasfilename(title=u'保存文件', filetypes=[("word文件", ".docx")],
                                             initialfile=file_name)
    print("保存位置目录", save_path)
    if save_path == "" or save_path is None:
        return
    # 导出到word
    downloadToWord(df_map, save_path, heading)
    showinfo("", "导出成功")


if __name__ == '__main__':
    gui = GuiTempldate()
    gui.showUI()
