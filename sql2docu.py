from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from pgsql.mypgsql import schema_info_s, Conn

conn = Conn("data", "postgres", "pwd", "192.168.1.1", "5432", "xhz")

# 输入表名
input_string = '''
'''

tables = input_string.strip().split('\n')
# 直接在列表推导式中处理每个表名，去除"_new"后缀
tables = [name[:-4] if name.endswith('_new') else name for name in tables]


# for i in tables:
#     print(str(i))


def write_to_word(result, filename):
    document = Document()
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # Define styles for table elements
    header_font_size = Pt(10)
    cell_font_size = Pt(10)

    # Loop through the result and create tables in the Word document
    for item in result:
        # print(f'item:{item}')
        table_name = item[0][0]
        table_comment = item[0][1]

        # Add an empty paragraph before each new table
        document.add_paragraph()

        # Create a new table for table name and comment
        table_header = document.add_table(rows=1, cols=1)
        table_header.width = Pt(500)
        table_header.style = 'Table Grid'

        # Add table name as a cell in the header table
        cell_name = table_header.cell(0, 0)
        cell_name.text = table_name if table_name else "表名"  # Use "表名" if table_name is None or empty
        cell_name.paragraphs[0].runs[0].font.size = header_font_size
        cell_name.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add table comment as a cell in the header table
        cell_comment = table_header.add_row().cells[0]
        cell_comment.text = table_comment.upper() if table_comment else "N/A"  # Use "N/A" if table_comment is None or empty
        cell_comment.paragraphs[0].runs[0].font.size = header_font_size
        cell_comment.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Create a new table for the actual data
        table_data = document.add_table(rows=1, cols=7)
        table_data.width = Pt(500)
        table_data.style = 'Table Grid'

        # Set header row style
        header_row = table_data.rows[0]
        header_row.cells[0].text = "字段名"
        header_row.cells[1].text = "注释"
        header_row.cells[2].text = "数据类型"
        header_row.cells[3].text = "长度"
        header_row.cells[4].text = "小数位数"
        header_row.cells[5].text = "是否允许为空"
        header_row.cells[6].text = "默认值"

        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].font.size = header_font_size
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Set cell style for the rest of the rows
        for row in item[1:]:
            new_row = table_data.add_row().cells
            for i, cell_data in enumerate(row[2:]):
                new_row[i].text = str(cell_data).upper()
                new_row[i].paragraphs[0].runs[0].font.size = cell_font_size
                new_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    document.save(filename)
    print("写入成功")


def write_to_excel(result, filename):
    wb = Workbook()
    ws = wb.active
    # 设置表头样式
    header_font = Font(bold=True)
    header_fill = PatternFill("solid", fgColor="92d050")  # 绿色背景色
    data_alignment = Alignment(horizontal='center', vertical='center')
    # 设置边框样式
    border = Border(left=Side(border_style='thin', color='000000'),
                    right=Side(border_style='thin', color='000000'),
                    top=Side(border_style='thin', color='000000'),
                    bottom=Side(border_style='thin', color='000000'))
    row_num = 0
    row_num_list = []
    for index, item in enumerate(result):
        if index == 0:
            row_num = 1
        else:
            length = len(result[index - 1])
            row_num += length + 4
            # 写入表注释
        row_num_list.append(row_num)
        # 写入表注释和表名
        ws.merge_cells(f"A{row_num}:G{row_num}")
        ws.merge_cells(f"A{row_num + 1}:G{row_num + 1}")
        for i in range(2):
            table_name_cell = ws.cell(row=row_num + i, column=1, value=item[0][i])
            table_name_cell.font = header_font
            table_name_cell.alignment = data_alignment
            table_name_cell.fill = header_fill
        header = ["字段名", "注释", "数据类型", "长度", "小数位数", "是否允许为空", "默认值"]
        ws.append(header)
        for cell in ws[ws.max_row]:
            cell.font = header_font
            cell.fill = header_fill  # 设置填充颜色
            cell.alignment = data_alignment
        # 写入数据剩下的数据
        for row in item:
            ws.append(row[2:])
    for row in ws.iter_rows():
        for cell in row:
            if cell.row + 1 not in row_num_list:
                cell.border = border
    wb.save(filename)
    print("写入成功")


line = '''
---------------------------------------------------------------------------
'''
result = schema_info_s(tables, conn)
for item in result:
    for row in item:
        print(f"row:{row}")
    # print(line)
print("执行写入函数")
filename = "excel_name"
docx_name = f"{filename}.docx"
xlsx_name = f"{filename}.xlsx"
write_to_excel(result, xlsx_name)

# write_to_word(result, filename)
